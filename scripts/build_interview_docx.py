from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "INTERVIEW_PREP.md"
OUTPUT = ROOT / "docs" / "AI_Course_Generator_Interview_Guide.docx"


def set_run_font(run, size=11, bold=False, italic=False, color="000000", font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0)
    r = p.add_run(text.strip())
    set_run_font(r, size=10, bold=bold)


def add_heading(doc, text, level):
    if level == 1:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=20, after=6)
        run = p.add_run(text)
        set_run_font(run, size=20, bold=False)
        return p
    if level == 2:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=18, after=6)
        run = p.add_run(text)
        set_run_font(run, size=16, bold=False)
        return p
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=16, after=4)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=False, color="434343")
    return p


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    segments = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for segment in segments:
        if not segment:
            continue
        if segment.startswith("**") and segment.endswith("**"):
            run = p.add_run(segment[2:-2])
            set_run_font(run, bold=True)
        elif segment.startswith("`") and segment.endswith("`"):
            run = p.add_run(segment[1:-1])
            set_run_font(run, font="Courier New", color="434343")
        else:
            run = p.add_run(segment)
            set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4)
    run = p.add_run(text)
    set_run_font(run)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4)
    run = p.add_run(text)
    set_run_font(run)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8)
    text = "\n".join(lines).strip()
    run = p.add_run(text)
    set_run_font(run, size=9.5, font="Courier New", color="434343")
    p.paragraph_format.left_indent = Inches(0.25)


def add_markdown_table(doc, rows):
    if len(rows) < 2:
        return
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) >= 2 and all(re.match(r"^-{3,}$", c.strip()) for c in parsed[1]):
        data = [parsed[0]] + parsed[2:]
    else:
        data = parsed
    if not data:
        return
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for r_idx, row in enumerate(data):
        for c_idx, cell_text in enumerate(row[: len(data[0])]):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, cell_text, bold=r_idx == 0)
            if r_idx == 0:
                set_cell_shading(cell, "F1F3F4")
    doc.add_paragraph()


def clean_markdown_text(text):
    text = text.replace("&", "&")
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    return text.strip()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("AI Course Generator")
    set_run_font(title_run, size=26, bold=False)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12)
    sub_run = subtitle.add_run("Interview Preparation Guide | MERN + Gemini AI Project")
    set_run_font(sub_run, size=11, color="555555")

    add_body_paragraph(
        doc,
        "Use this document to prepare project explanations, architecture answers, feature walkthroughs, and technical interview responses.",
    )

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    table_rows = []
    skip_first_title = True

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            flush_table()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            if line.strip().startswith(("flowchart", "sequenceDiagram")):
                code_lines.append(line.strip())
            elif code_lines or line.strip():
                code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            table_rows.append(line)
            continue
        else:
            flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            if skip_first_title:
                skip_first_title = False
                continue
            add_heading(doc, clean_markdown_text(line[2:]), 1)
        elif line.startswith("## "):
            add_heading(doc, clean_markdown_text(line[3:]), 1)
        elif line.startswith("### "):
            add_heading(doc, clean_markdown_text(line[4:]), 2)
        elif line.startswith("#### "):
            add_heading(doc, clean_markdown_text(line[5:]), 3)
        elif line.startswith("- "):
            add_bullet(doc, clean_markdown_text(line[2:]))
        elif re.match(r"^\d+\.\s+", line):
            add_numbered(doc, clean_markdown_text(re.sub(r"^\d+\.\s+", "", line)))
        elif line.startswith("> "):
            p = add_body_paragraph(doc, clean_markdown_text(line[2:]))
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor.from_string("555555")
        else:
            add_body_paragraph(doc, clean_markdown_text(line))

    flush_table()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
